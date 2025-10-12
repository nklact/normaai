"""
DOCX Generation Utility
Creates professional .docx files from contract text.
"""

import os
import uuid
from datetime import datetime
from typing import Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


class ContractDocxGenerator:
    """Generates professional .docx contracts from text."""

    def __init__(self, temp_dir: str = "/tmp/contracts"):
        """
        Initialize the generator.

        Args:
            temp_dir: Directory to store generated contracts
        """
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)

    def generate_contract(
        self,
        contract_content: str,
        contract_type: str = "Ugovor"
    ) -> Tuple[str, str, str]:
        """
        Generate a .docx file from contract content.

        Args:
            contract_content: The contract text
            contract_type: Type of contract for title

        Returns:
            Tuple of (file_id, filepath, filename)
        """
        # Create document
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Add title
        title = doc.add_heading(contract_type.upper(), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.name = 'Times New Roman'

        # Add spacing after title
        doc.add_paragraph()

        # Process and add contract content
        self._add_contract_content(doc, contract_content)

        # Add footer with generation date
        self._add_footer(doc)

        # Generate unique file ID
        file_id = str(uuid.uuid4())

        # Create filename
        timestamp = datetime.now().strftime("%Y-%m-%d")
        safe_type = contract_type.replace(" ", "_").replace("/", "-")
        filename = f"{safe_type}_{timestamp}.docx"

        # Save file
        filepath = os.path.join(self.temp_dir, f"{file_id}.docx")
        doc.save(filepath)

        return file_id, filepath, filename

    def _add_contract_content(self, doc: Document, content: str):
        """Add formatted contract content to document."""
        # Split into sections by double newlines
        sections = content.split('\n\n')

        for section in sections:
            section = section.strip()
            if not section:
                continue

            # Check if it's a heading (e.g., "Član 1.", "I. PREDMET UGOVORA")
            if self._is_heading(section):
                para = doc.add_paragraph(section)
                para.style = 'Heading 2'
                para_format = para.paragraph_format
                para_format.space_before = Pt(12)
                para_format.space_after = Pt(6)
                para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True
            else:
                # Regular paragraph
                para = doc.add_paragraph(section)
                para_format = para.paragraph_format
                para_format.space_after = Pt(6)
                para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                para_format.first_line_indent = Inches(0.5)
                para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    def _is_heading(self, text: str) -> bool:
        """Determine if text should be formatted as a heading."""
        text_lower = text.lower()

        # Common Serbian contract headings
        heading_indicators = [
            text.startswith("član"),
            text.startswith("stav"),
            text.startswith("i."),
            text.startswith("ii."),
            text.startswith("iii."),
            text.startswith("iv."),
            text.startswith("v."),
            "predmet ugovora" in text_lower,
            "osnovne odredbe" in text_lower,
            "prava i obaveze" in text_lower,
            "završne odredbe" in text_lower,
            len(text) < 80 and text.endswith(":"),
        ]

        return any(heading_indicators)

    def _add_footer(self, doc: Document):
        """Add footer with generation info."""
        section = doc.sections[0]
        footer = section.footer

        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = f"Generisano uz pomoć Norma AI • {datetime.now().strftime('%d.%m.%Y.')}"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in footer_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.name = 'Times New Roman'

    def get_file_path(self, file_id: str) -> str:
        """Get the file path for a given file ID."""
        return os.path.join(self.temp_dir, f"{file_id}.docx")

    def file_exists(self, file_id: str) -> bool:
        """Check if a file exists."""
        return os.path.exists(self.get_file_path(file_id))

    def delete_file(self, file_id: str) -> bool:
        """
        Delete a contract file.

        Returns:
            True if deleted, False if file didn't exist
        """
        filepath = self.get_file_path(file_id)
        if os.path.exists(filepath):
            try:
                os.remove(filepath)
                return True
            except Exception as e:
                print(f"Error deleting file {filepath}: {e}")
                return False
        return False


# Example usage
if __name__ == "__main__":
    sample_contract = """
    UGOVOR O RADU NA NEODREĐENO VREME

    Zaključen dana _____________ između:

    1. TECH DOO, sa sedištem u Beogradu
    (u daljem tekstu: Poslodavac)

    i

    2. Marko Marković, sa prebivalištem u ____________
    (u daljem tekstu: Zaposleni)

    Član 1.
    Poslodavac zaključuje sa Zaposlenim ugovor o radu na neodređeno vreme, počev od 01.11.2025. godine.

    Član 2.
    Zaposleni će obavljati poslove Software Developera u skladu sa Pravilnikom o sistematizaciji radnih mesta.

    Član 3.
    Zaposleni ima pravo na neto zaradu u iznosu od 150.000,00 RSD mesečno, koja se isplaćuje do 15-og u mesecu za prethodni mesec.

    I. RADNO VREME

    Član 4.
    Puno radno vreme je 40 časova nedeljno, sa rasporedom od ponedeljka do petka, od 09:00 do 17:00 časova.

    II. ZAVRŠNE ODREDBE

    Član 10.
    Ugovor je sačinjen u 3 (tri) istovetna primerka, po jedan za svaku ugovornu stranu i jedan za nadležni organ.
    """

    generator = ContractDocxGenerator(temp_dir="./test_contracts")
    file_id, filepath, filename = generator.generate_contract(
        sample_contract,
        "Ugovor o radu"
    )

    print(f"Generated contract:")
    print(f"  File ID: {file_id}")
    print(f"  Filepath: {filepath}")
    print(f"  Filename: {filename}")
    print(f"  Exists: {generator.file_exists(file_id)}")
